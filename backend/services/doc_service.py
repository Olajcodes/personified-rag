# This handles word document creation

import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocService:
    @staticmethod
    def create_docx_from_text(text: str):
        """
        Docx generator that matches the reference image style:
        - Main Name Header (#): Centered
        - Contact Line (with '|'): Centered
        - Section Headers (##, ###): Left Aligned
        - Body Text: Justified
        """
        doc = Document()
        
        # Set Default Font
        style = doc.styles['Normal']
        style.font.name = 'Cambria'
        style.font.size = Pt(12)

        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # --- HEADERS ---
            if line.startswith('# ') or line.startswith('## ') or line.startswith('### '):
                clean_text = line.replace('#', '').strip()
                level = 1
                if line.startswith('### '): level = 3
                elif line.startswith('## '): level = 2
                
                heading = doc.add_heading(clean_text, level=level)
                
                # LOGIC: Center ONLY the main title (Level 1), Left align others
                if level == 1:
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Styling for H1 & H2 (Dark Blue)
                if level <= 2:
                    run = heading.runs[0]
                    run.font.color.rgb = RGBColor(0, 51, 102) 
                    run.font.name = 'Cambria'
                    run.font.bold = True
                
            # --- BULLET POINTS (Left Aligned) ---
            elif line.startswith('- ') or line.startswith('* '):
                clean_text = line[2:] 
                p = doc.add_paragraph(style='List Bullet')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                parts = clean_text.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1: 
                        run.font.bold = True

            # --- STANDARD PARAGRAPHS ---
            else:
                p = doc.add_paragraph()
                
                # LOGIC: Center Contact Info line, Justify everything else
                if '|' in line:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.font.bold = True
                
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream