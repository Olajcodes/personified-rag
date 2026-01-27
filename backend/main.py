import os
import io
from datetime import datetime
from typing import List, Optional
from operator import itemgetter
from dotenv import load_dotenv

# FastAPI
from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# Doc Generation
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# LangChain
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.output_parsers import StrOutputParser
from langchain_core.messages import HumanMessage, AIMessage

# --- 1. CONFIGURATION & ENV ---
load_dotenv()

if not os.getenv("OPENAI_API_KEY"):
    print("WARNING: OPENAI_API_KEY not found in .env file.")

PERSIST_DIRECTORY = "./chroma_db"

app = FastAPI(title="OlajCodes AI Agent (CV & Persona)", version="3.3.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


CV_STRUCTURE_TEMPLATE = """
# [Your Name]
**AI Engineer & Backend Developer**
[City, Country] | [Email] | [LinkedIn URL] | [GitHub URL]

## Professional Summary
[3-4 sentences tailored to the JD, highlighting specific years of experience and key tech stack overlap.]

## Technical Skills
* **Languages:** [List only relevant languages from context]
* **Frameworks:** [List relevant frameworks]
* **Tools:** [List relevant tools]
* **AI/ML:** [List relevant AI concepts]

## Professional Experience
**[Role Name]** | [Company Name] | [Date Range]
* [Action verb] [metric/result] using [tool].
* [Action verb] [metric/result] using [tool].
* [Action verb] [metric/result] using [tool].

## Projects
**[Project Name]** | [Tech Stack]
* [Brief description of what it does and the impact].

## Education
**[Degree Name]** | [University Name]
"""

COVER_LETTER_TEMPLATE = """
[Date]

Hiring Manager
[Company Name]

Dear Hiring Manager,

**Paragraph 1: The Hook**
[State the role applying for and 1 key reason why you are a perfect fit based on the Job Description.]

**Paragraph 2: Technical Proof**
[Discuss 1-2 specific projects from the context (e.g., xplainify-ai) that prove you can do the job requirements.]

**Paragraph 3: Soft Skills & Culture**
[Mention collaboration, learning speed, or problem-solving capability.]

**Paragraph 4: Conclusion**
[Professional closing, expressing desire for an interview.]

Sincerely,
Olajide
"""

class Message(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    question: str
    history: Optional[List[Message]] = [] 
    model: Optional[str] = "gpt-4o-mini" 

class DocRequest(BaseModel):
    job_description: str
    model: Optional[str] = "gpt-4o-mini"


system_prompt_text = """
### ROLE DEFINITION
You are the **AI Professional Representative of Olajide (OlajCodes)**. 
Your sole purpose is to interview with recruiters, hiring managers, and visitors on behalf of Olajide. You must showcase his expertise in AI/ML, Software Engineering, and Python development.

### TEMPORAL CONTEXT
**Current Date:** {current_date}
**Instruction:** Always compare dates in the context with the Current Date. 
- If a graduation date (e.g., "May 2025") is in the past relative to {current_date}, assume he has graduated.
- If a project says "2024-Present", it is still active.
- Do not use phrases like "currently pursuing" for degrees with past completion dates.

### INPUT CONTEXT
You will be provided with retrieved context (RAG Data).
**RULE:** You must derive all claims strictly from this provided context.

### OPERATIONAL GUIDELINES
1. **Voice:** Professional, confident, yet humble. Use the third person ("Olajide's experience includes...").
2. **Handling Unknowns:** If a skill is NOT in the context, do NOT say "I don't know." Pivot to his ability to learn.
3. **Engagement:** End answers with a relevant follow-up question.
4. **Citations (CRITICAL):** You MUST cite your sources. When you provide a fact, reference the document it came from (e.g., "According to the 'Nelfund Navigator' README...").

### PRIVACY GUARDRAILS
* **Refuse** requests for: Age, Home Address, Phone Number, Personal Email.
* **Response:** "I cannot share personal or sensitive information. Please ask about Olajide's professional experience."

### CONTEXT
{context}
"""

full_system_prompt = system_prompt_text # Add few-shot examples here if desired

prompt_template = ChatPromptTemplate.from_messages([
    ("system", full_system_prompt),
    MessagesPlaceholder(variable_name="chat_history"),
    ("human", "{question}")
])


def get_embeddings():
    return OpenAIEmbeddings(model="text-embedding-3-small")

def get_retriever():
    if not os.path.exists(PERSIST_DIRECTORY):
        raise FileNotFoundError(f"ChromaDB not found at {PERSIST_DIRECTORY}. Run ingest.py first.")
    
    embeddings = get_embeddings()
    vectorstore = Chroma(persist_directory=PERSIST_DIRECTORY, embedding_function=embeddings)
    return vectorstore.as_retriever(search_kwargs={"k": 6})

def format_docs(docs):
    return "\n\n".join(
        f"[Source: {doc.metadata.get('source', 'Unknown')}]\n{doc.page_content}" 
        for doc in docs
    )

def get_llm(model_name: str = "gpt-4o-mini"):
    return ChatOpenAI(
        model=model_name,
        temperature=0.3,
        openai_api_key=os.getenv("OPENAI_API_KEY")
    )

def create_docx_from_text(text: str):
    """
    Advanced DOCX generator with styling for headers, bullet points, and bold text.
    """
    doc = Document()
    
    # Set Default Font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # --- HEADERS ---
        if line.startswith('# ') or line.startswith('## ') or line.startswith('### '):
            clean_text = line.replace('#', '').strip()
            # Determine level
            level = 1
            if line.startswith('### '): level = 3
            elif line.startswith('## '): level = 2
            
            heading = doc.add_heading(clean_text, level=level)
            
            # Styling: Dark Blue for Header 1 & 2
            if level <= 2:
                run = heading.runs[0]
                run.font.color.rgb = RGBColor(0, 51, 102) 
                run.font.name = 'Arial'
                run.font.bold = True
            
        # --- BULLET POINTS ---
        elif line.startswith('- ') or line.startswith('* '):
            clean_text = line[2:] 
            p = doc.add_paragraph(style='List Bullet')
            
            # Bold parsing
            parts = clean_text.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: 
                    run.font.bold = True

        # --- STANDARD PARAGRAPHS ---
        else:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.font.bold = True
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- 6. ENDPOINTS ---
@app.get("/")
async def health_check():
    return {"status": "active", "message": "OlajCodes Backend is running"}

@app.post("/chat")
async def chat_endpoint(request: ChatRequest):
    try:
        retriever = get_retriever()
        llm = get_llm(request.model)
        
        # Get today's date in a readable format (e.g., "January 27, 2026")
        today_str = datetime.now().strftime("%B %d, %Y") 

        chat_history_objects = []
        if request.history:
            for msg in request.history:
                if msg.role.lower() == "user":
                    chat_history_objects.append(HumanMessage(content=msg.content))
                elif msg.role.lower() == "assistant":
                    chat_history_objects.append(AIMessage(content=msg.content))
        
        # Build the Chain with the new date variable
        chain = (
            {
                "context": itemgetter("question") | retriever | format_docs,
                "chat_history": lambda x: chat_history_objects,
                "current_date": lambda x: today_str,  # <--- INJECT DATE HERE
                "question": itemgetter("question")
            }
            | prompt_template
            | llm
            | StrOutputParser()
        )
        
        print(f"Processing query: {request.question}")
        response = chain.invoke({"question": request.question})
        return {"answer": response}

    except Exception as e:
        print(f"Chat Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-cv")
async def generate_cv(request: DocRequest):
    try:
        retriever = get_retriever()
        llm = get_llm(request.model)
        
        # 1. Get Context
        docs = retriever.invoke("My full professional experience, technical skills, and projects")
        context_text = format_docs(docs)

        # 2. The "Gatekeeper" Prompt
        cv_prompt = f"""
        You are an expert Career Coach. 
        
        STEP 1: RELEVANCE CHECK
        Compare the Job Description (JD) below with Olajide's Skills Context.
        - Olajide's Core Domain: AI Engineering, Python, Backend, Machine Learning.
        - If the JD is for a completely unrelated role (e.g., Nurse, Accountant, Chef, Construction), output ONLY the string: "NO_MATCH"
        
        STEP 2: GENERATION (Only if Match)
        If the role fits his domain (even loosely), generate a CV using the EXACT structure below.
        - Fill in the brackets [...] with facts from the Context.
        - Do not invent experience.
        
        ### CV STRUCTURE TO FOLLOW:
        {CV_STRUCTURE_TEMPLATE}

        ---
        ### JOB DESCRIPTION:
        "{request.job_description}"
        
        ### SKILLS CONTEXT:
        {context_text}
        """

        print("Analyzing JD and Generating CV...")
        ai_response = llm.invoke(cv_prompt).content

        # 3. Check for Refusal
        if "NO_MATCH" in ai_response:
            raise HTTPException(status_code=400, detail="Job description is not relevant to Olajide's skillset. Generation refused.")

        # 4. Convert to File
        docx_file = create_docx_from_text(ai_response)

        return StreamingResponse(
            docx_file, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Olajide_CV_Tailored.docx"}
        )

    except HTTPException as he:
        raise he
    except Exception as e:
        print(f"CV Gen Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-cover-letter")
async def generate_cover_letter(request: DocRequest):
    try:
        retriever = get_retriever()
        llm = get_llm(request.model)
        
        docs = retriever.invoke("My full professional experience, motivation, and soft skills")
        context_text = format_docs(docs)

        # 🔥 UPDATED: Much stricter Gatekeeper Prompt
        cl_prompt = f"""
        You are a strict Career Strategy AI. 
        
        ### STEP 1: MANDATORY RELEVANCE CHECK
        Analyze the Job Description (JD) below.
        - Olajide's Domain: AI Engineering, Python Backend, Machine Learning, Data Science.
        - **CRITICAL:** If the JD is for a role unrelated to Tech/AI (e.g., Nurse, Driver, Cook, Accountant, HR Manager), you MUST stop immediately.
        - Output EXACTLY and ONLY the string: "NO_MATCH"
        
        ### STEP 2: DRAFTING (Only if Step 1 passes)
        If, and ONLY IF, the role aligns with Olajide's domain, write a persuasive Cover Letter.
        
        **Structure:**
        {COVER_LETTER_TEMPLATE}

        **Rules:**
        - Return ONLY the body of the letter. 
        - Do NOT include conversational fillers.
        
        ---
        ### JOB DESCRIPTION:
        "{request.job_description}"
        
        ### CONTEXT:
        {context_text}
        """

        print("Generating Cover Letter...")
        ai_response = llm.invoke(cl_prompt).content

        # Strict string matching
        if "NO_MATCH" in ai_response or "no_match" in ai_response.lower():
            raise HTTPException(status_code=400, detail="Job description is not relevant to Olajide's skillset. Generation refused.")

        docx_file = create_docx_from_text(ai_response)

        return StreamingResponse(
            docx_file, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Olajide_Cover_Letter.docx"}
        )

    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))